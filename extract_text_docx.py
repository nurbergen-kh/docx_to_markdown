from docx import Document

def extract_tables(file_path):
    """Extract all text from DOCX using python-docx"""
    doc = Document(file_path)
    tables = []
    
    for table in doc.tables:
        rows = []
        for row in table.rows:
            columns = []
            text = set()
            for cell in row.cells:
                if cell.text.strip():
                    if cell.text not in text:
                        columns.append(cell.text)
                        text.add(cell.text)
            if columns:
                rows.append(columns)

        tables.append(rows)
    
    return tables


def extract_paragraphs(file_path):
    """Extract all text from DOCX using python-docx"""
    doc = Document(file_path)
    paragraphs = []
    
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)
    
    return paragraphs


if __name__ == '__main__':
    import sys
    if len(sys.argv) != 2:
        print("Usage: python extract_text_docx.py input.docx")
        sys.exit(1)
    
    # tables = extract_tables(sys.argv[1])
    # paragraphs = extract_paragraphs(sys.argv[1])
